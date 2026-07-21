import io
import logging
from typing import List
import pypdf
import docx

from sqlalchemy.ext.asyncio import AsyncSession
from app.repositories.document_repository import DocumentRepository
from app.storage.base import StorageProvider
from app.document_viewer.models import DocumentViewerInfo, DocumentPageData, DocumentSearchMatch

logger = logging.getLogger("NEXO.ViewerService")


class DocumentViewerService:
    """Handles text extraction and page parsing for explainable viewer."""

    def __init__(self, session: AsyncSession, storage: StorageProvider) -> None:
        self.session = session
        self.storage = storage
        self.repo = DocumentRepository(session)

    async def get_viewer_info(self, document_id: str) -> DocumentViewerInfo:
        document = await self.repo.get_by_id(document_id)
        if not document:
            raise ValueError("Document not found")

        # Extract page count
        total_pages = 1
        try:
            file_bytes = await self.storage.download(document.stored_filename)
            if document.original_filename.endswith(".pdf"):
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                total_pages = len(reader.pages)
            elif document.original_filename.endswith(".docx"):
                doc = docx.Document(io.BytesIO(file_bytes))
                total_pages = max(1, (len(doc.paragraphs) + 4) // 5)
        except Exception as e:
            logger.warning(f"Error calculating page count: {str(e)}. Fallback to 1 page.")
            total_pages = 1

        return DocumentViewerInfo(
            document_id=document.id,
            title=document.title,
            mime_type=document.mime_type or "application/octet-stream",
            total_pages=total_pages,
            status=document.status,
            file_size=document.file_size or len(file_bytes)
        )

    async def get_page_content(self, document_id: str, page_number: int) -> DocumentPageData:
        document = await self.repo.get_by_id(document_id)
        if not document:
            raise ValueError("Document not found")

        text = ""
        file_bytes = await self.storage.download(document.stored_filename)
        
        try:
            if document.original_filename.endswith(".pdf"):
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                if 1 <= page_number <= len(reader.pages):
                    text = reader.pages[page_number - 1].extract_text() or ""
            elif document.original_filename.endswith(".docx"):
                doc = docx.Document(io.BytesIO(file_bytes))
                para_start = (page_number - 1) * 5
                para_end = para_start + 5
                text_parts = [p.text for p in doc.paragraphs[para_start:para_end]]
                text = "\n\n".join(text_parts)
            else:
                # Text/Markdown files
                text = file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.warning(f"Error extracting page text: {str(e)}. Using fallback text.")
            text = f"Fallback page content for page {page_number}. Error: {str(e)}"

        return DocumentPageData(
            page_number=page_number,
            text_content=text,
            highlights=[]
        )

    async def search_document(self, document_id: str, query: str) -> List[DocumentSearchMatch]:
        document = await self.repo.get_by_id(document_id)
        if not document:
            raise ValueError("Document not found")

        matches = []
        info = await self.get_viewer_info(document_id)
        
        for p in range(1, info.total_pages + 1):
            page_data = await self.get_page_content(document_id, p)
            content = page_data.text_content
            if query.lower() in content.lower():
                # Extract snippet
                idx = content.lower().find(query.lower())
                start = max(0, idx - 40)
                end = min(len(content), idx + len(query) + 40)
                snippet = "..." + content[start:end] + "..."
                matches.append(
                    DocumentSearchMatch(
                        page_number=p,
                        snippet=snippet,
                        text_match=content[idx:idx+len(query)]
                    )
                )

        return matches
